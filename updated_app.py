import streamlit as st
import json
import requests
import io
import os
from dotenv import load_dotenv
from openai import OpenAI
from pypdf import PdfReader
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# 1. Load Config
load_dotenv()
st.set_page_config(page_title="Resume Architect AI", page_icon="📝", layout="wide")

# 2. CSS for Professional Look
st.markdown("""
<style>
    .stButton>button { width: 100%; border-radius: 5px; height: 3em; background-color: #2b6cb0; color: white; font-weight: bold; }
</style>
""", unsafe_allow_html=True)

# 3. API Setup
api_key = st.secrets.get("OPENAI_API_KEY")
if not api_key:
    # Fallback for local testing
    api_key = os.getenv("OPENAI_API_KEY")

if not api_key:
    st.error("🚨 API Key Missing.")
    st.stop()

client = OpenAI(api_key=api_key)

# --- HELPER FUNCTIONS ---

def fetch_jd_from_url(url):
    """
    Fetches text from a URL using Jina Reader.
    """
    try:
        jina_url = f"https://r.jina.ai/{url}"
        response = requests.get(jina_url, timeout=10)
        return response.text
    except Exception as e:
        return None

def create_docx(data):
    """
    Generates a professional Word Doc from the AI JSON data.
    """
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(10)

    # 1. Header (Name & Contact)
    h1 = doc.add_paragraph()
    h1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = h1.add_run(data.get('name', 'Candidate Name'))
    run.bold = True
    run.font.size = Pt(24)
    run.font.color.rgb = RGBColor(0, 51, 102) # Dark Blue

    contact = doc.add_paragraph()
    contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact.add_run(data.get('contact_info', 'Location | Email | Phone | LinkedIn'))

    # 2. Professional Summary
    if data.get('summary'):
        doc.add_heading('PROFESSIONAL SUMMARY', level=1)
        doc.add_paragraph(data['summary'])

    # 3. Technical Skills
    if data.get('skills'):
        doc.add_heading('TECHNICAL SKILLS', level=1)
        if isinstance(data['skills'], dict):
            for category, items in data['skills'].items():
                p = doc.add_paragraph()
                p.add_run(f"{category}: ").bold = True
                p.add_run(items)
        else:
            doc.add_paragraph(str(data['skills']))

    # 4. Professional Experience
    if data.get('experience'):
        doc.add_heading('PROFESSIONAL EXPERIENCE', level=1)
        for job in data['experience']:
            # Role line
            p = doc.add_paragraph()
            run_role = p.add_run(f"{job.get('role', 'Role')} | {job.get('company', 'Company')}")
            run_role.bold = True
            run_role.font.size = Pt(11)
            
            # Dates
            if job.get('dates'):
                p.add_run(f"  ({job['dates']})").italic = True
            
            # Bullets
            for bullet in job.get('bullets', []):
                doc.add_paragraph(bullet, style='List Bullet')

    # 5. Education
    if data.get('education'):
        doc.add_heading('EDUCATION', level=1)
        for edu in data['education']:
            doc.add_paragraph(f"{edu.get('degree', '')} - {edu.get('school', '')}")

    # Save to memory stream
    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    return file_stream

def analyze_fit(resume_text, jd_text):
    """
    Strict Analysis with Consistency Locks.
    """
    prompt = f"""
    Act as a strict ATS. Compare the RESUME vs JD.
    
    SCORING RULES:
    - 100%: Perfect match (All hard skills + exact experience).
    - <60%: Missing critical hard skills.
    - Ignore Soft Skills. Focus on Tools, Languages, Frameworks.

    Return JSON:
    {{
        "match_percentage": Integer (0-100),
        "missing_keywords": ["List", "of", "missing", "hard", "skills"],
        "summary_critique": "1 sentence critique."
    }}
    
    JD: {jd_text[:4000]}
    RESUME: {resume_text[:4000]}
    """
    response = client.chat.completions.create(
        model="gpt-4o-mini",
        messages=[{"role": "user", "content": prompt}],
        temperature=0.0, # Lock creativity
        seed=42,         # Lock randomness
        response_format={"type": "json_object"}
    )
    return json.loads(response.choices[0].message.content)

def rewrite_resume_to_json(resume_text, jd_text, gap_data):
    """
    Rewrites resume to JSON structure for Word Doc generation.
    Uses Google XYZ formula.
    """
    prompt = f"""
    You are a Top-Tier Resume Writer. 
    Rewrite the RESUME to target the JD, using the 'Google XYZ Formula' for bullet points.
    
    Formula: "Accomplished [X] as measured by [Y], by doing [Z]".
    
    INSTRUCTIONS:
    1. Integrate missing keywords: {gap_data['missing_keywords']}
    2. Maintain the user's truth but quantify results.
    3. Return strictly JSON matching the structure below.

    JSON Structure:
    {{
        "name": "Candidate Name",
        "contact_info": "City | Email | Phone",
        "summary": "Revised Summary",
        "skills": {{"Category": "List of skills"}},
        "experience": [
            {{
                "role": "Job Title",
                "company": "Company",
                "dates": "Date - Date",
                "bullets": ["XYZ Bullet 1", "XYZ Bullet 2"]
            }}
        ],
        "education": [{{ "degree": "Degree", "school": "University" }}]
    }}

    JD: {jd_text[:4000]}
    RESUME: {resume_text[:4000]}
    """
    response = client.chat.completions.create(
        model="gpt-4o",
        messages=[{"role": "user", "content": prompt}],
        temperature=0.0, # Consistent Output
        seed=42,
        response_format={"type": "json_object"}
    )
    return json.loads(response.choices[0].message.content)

# --- MAIN UI ---

st.title("🚀 Resume Architect AI")
st.markdown("Upload your resume and a JD. We will **Audit** and **Rewrite** it into an **Editable Word Doc**.")

# 1. Inputs
col1, col2 = st.columns(2)

with col1:
    st.subheader("1. The Job")
    tab_text, tab_link = st.tabs(["Paste Text", "Paste Link"])
    
    with tab_text:
        jd_text_input = st.text_area("Paste Job Description", height=200)
    
    with tab_link:
        jd_url = st.text_input("Paste LinkedIn/Indeed URL")
        fetched_jd = None
        if jd_url:
            with st.spinner("Fetching Job Description..."):
                fetched_jd = fetch_jd_from_url(jd_url)
                if fetched_jd:
                    st.success("✅ JD Loaded from Link")
                    with st.expander("View Fetched Text"):
                        st.text(fetched_jd[:500] + "...")
                else:
                    st.error("Could not fetch link.")

with col2:
    st.subheader("2. Your Resume")
    uploaded_file = st.file_uploader("Upload PDF", type="pdf")

# 2. Logic
if st.button("🚀 Analyze & Rewrite"):
    # Determine JD Source
    final_jd = jd_text_input if jd_text_input else fetched_jd
    
    if not final_jd or not uploaded_file:
        st.warning("Please provide both a JD and a Resume.")
    else:
        # A. Read PDF
        with st.spinner("Reading PDF..."):
            reader = PdfReader(uploaded_file)
            original_text = ""
            for page in reader.pages:
                original_text += page.extract_text() + "\n"

        # B. Analyze (Original)
        with st.spinner("Analyzing Fit..."):
            original_analysis = analyze_fit(original_text, final_jd)

        # C. Rewrite (To JSON)
        with st.spinner("Rewriting using 'Google XYZ Formula'..."):
            resume_json = rewrite_resume_to_json(original_text, final_jd, original_analysis)
            
            # Generate Docx
            docx_data = create_docx(resume_json)

        # D. Analyze (New Version) - Fairness Check
        with st.spinner("Verifying New Score..."):
            # Convert JSON back to text for fair comparison
            new_text_str = json.dumps(resume_json)
            new_analysis = analyze_fit(new_text_str, final_jd)

        # --- RESULTS ---
        st.divider()
        st.subheader("📊 Optimization Report")

        c1, c2, c3 = st.columns(3)
        with c1:
            st.metric("Original Score", f"{original_analysis['match_percentage']}%")
        with c2:
            st.metric("Optimized Score", f"{new_analysis['match_percentage']}%")
        with c3:
            imp = new_analysis['match_percentage'] - original_analysis['match_percentage']
            st.metric("Improvement", f"+{imp}%")

        # Gap Analysis
        col_a, col_b = st.columns(2)
        with col_a:
            st.write("### 🔴 Identified Gaps")
            st.warning(f"Missing Keywords: {', '.join(original_analysis['missing_keywords'])}")
        with col_b:
            st.write("### 🟢 Optimization Strategy")
            st.success("Applied 'Google XYZ' formula and integrated missing hard skills.")

        # Download
        st.divider()
        st.subheader("📥 Download Editable Resume")
        st.markdown("Your resume has been converted to an **Editable Word Document**.")
        
        st.download_button(
            label="Download .docx File",
            data=docx_data,
            file_name="Optimized_Resume.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
