import streamlit as st
import json
from openai import OpenAI
from pypdf import PdfReader
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io

# --- CONFIGURATION ---
st.set_page_config(page_title="Salesforce Resume Architect", page_icon="☁️", layout="wide")

api_key = st.secrets.get("OPENAI_API_KEY")
if not api_key:
    st.error("🚨 API Key Missing. Please check Streamlit Secrets.")
    st.stop()

client = OpenAI(api_key=api_key)

# --- HELPER: DOCX GENERATOR ---
def create_docx(data):
    """
    Builds a professional resume using python-docx based on JSON data.
    """
    doc = Document()
    
    # 1. Name & Contact (Center Aligned)
    h1 = doc.add_paragraph()
    h1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = h1.add_run(data.get('name', 'Candidate Name'))
    run.bold = True
    run.font.size = Pt(24)
    run.font.name = 'Arial'
    run.font.color.rgb = RGBColor(0, 112, 210) # Salesforce Blue

    contact = doc.add_paragraph()
    contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact.add_run(data.get('contact_info', 'Location | Email | Phone'))

    # 2. Summary
    doc.add_heading('PROFESSIONAL SUMMARY', level=1)
    doc.add_paragraph(data.get('summary', ''))

    # 3. Skills
    doc.add_heading('TECHNICAL SKILLS', level=1)
    # Create a nice looking "Key: Value" format
    skills = data.get('skills', {})
    for category, items in skills.items():
        p = doc.add_paragraph()
        run_cat = p.add_run(f"{category}: ")
        run_cat.bold = True
        p.add_run(items)

    # 4. Experience
    doc.add_heading('PROFESSIONAL EXPERIENCE', level=1)
    for job in data.get('experience', []):
        # Job Title Line
        p = doc.add_paragraph()
        run_title = p.add_run(f"{job['role']} | {job['company']}")
        run_title.bold = True
        run_title.font.size = Pt(11)
        
        # Dates (Right aligned technically hard in basic docx, so we append)
        p.add_run(f"  ({job['dates']})").italic = True
        
        # Bullets
        for bullet in job['bullets']:
            doc.add_paragraph(bullet, style='List Bullet')

    # 5. Education
    if data.get('education'):
        doc.add_heading('EDUCATION & CERTIFICATIONS', level=1)
        for edu in data.get('education', []):
            doc.add_paragraph(f"{edu['degree']} - {edu['school']}")

    # Save to memory
    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    return file_stream

# --- AI LOGIC ---

def analyze_gap(resume_text, jd_text):
    prompt = f"""
    You are a Salesforce Recruiter. Analyze RESUME vs JD.
    Return JSON:
    {{
        "score": 0,
        "missing_skills": ["skill1", "skill2"],
        "advice": "Strategic advice string"
    }}
    RESUME: {resume_text[:3000]}
    JD: {jd_text[:3000]}
    """
    response = client.chat.completions.create(
        model="gpt-4o-mini",
        messages=[{"role": "user", "content": prompt}],
        response_format={"type": "json_object"}
    )
    return json.loads(response.choices[0].message.content)

def generate_resume_json(resume_text, jd_text, gap_data):
    """
    Returns STRUCTURED JSON data instead of raw text/HTML.
    """
    prompt = f"""
    You are a Resume Writer. Rewrite the resume to target the JD.
    Integrate missing skills: {gap_data['missing_skills']}.
    Use STAR method for bullets.
    
    Return STRICT JSON with this structure:
    {{
        "name": "Candidate Name",
        "contact_info": "City, State | Email | Phone | LinkedIn",
        "summary": "Professional summary...",
        "skills": {{
            "Salesforce": "Apex, LWC, Flow...",
            "Tools": "Jira, Git..."
        }},
        "experience": [
            {{
                "role": "Job Title",
                "company": "Company Name",
                "dates": "Jan 2020 - Present",
                "bullets": ["Bullet 1", "Bullet 2"]
            }}
        ],
        "education": [
            {{"degree": "BS Computer Science", "school": "University Name"}}
        ]
    }}
    
    RESUME: {resume_text[:3000]}
    JD: {jd_text[:3000]}
    """
    response = client.chat.completions.create(
        model="gpt-4o",
        messages=[{"role": "user", "content": prompt}],
        response_format={"type": "json_object"}
    )
    return json.loads(response.choices[0].message.content)

def generate_cover_letter(resume_text, jd_text):
    prompt = f"""
    Write a short, punchy Cover Letter.
    RESUME: {resume_text[:3000]}
    JD: {jd_text[:3000]}
    """
    response = client.chat.completions.create(
        model="gpt-4o",
        messages=[{"role": "user", "content": prompt}]
    )
    return response.choices[0].message.content

# --- UI ---
st.title("☁️ Salesforce Resume Architect")
st.markdown("Upload your resume. We will generate an **Editable Word Doc** optimized for the job.")

col1, col2 = st.columns(2)
with col1:
    jd_input = st.text_area("Paste Job Description", height=200)
with col2:
    uploaded_file = st.file_uploader("Upload Resume (PDF)", type="pdf")

if 'processed' not in st.session_state:
    st.session_state.processed = False
    st.session_state.docx_data = None
    st.session_state.cover_letter = None
    st.session_state.analysis = None

if st.button("🚀 Analyze & Rewrite"):
    if not jd_input or not uploaded_file:
        st.warning("Please provide both inputs!")
    else:
        reader = PdfReader(uploaded_file)
        text = "".join([page.extract_text() for page in reader.pages])
        
        with st.spinner("Analyzing Gaps..."):
            st.session_state.analysis = analyze_gap(text, jd_input)
            
        with st.spinner("Structuring New Resume (GPT-4o)..."):
            # 1. Get JSON Data
            resume_json = generate_resume_json(text, jd_input, st.session_state.analysis)
            
            # 2. Build Word Doc
            st.session_state.docx_data = create_docx(resume_json)
            
        with st.spinner("Drafting Cover Letter..."):
            st.session_state.cover_letter = generate_cover_letter(text, jd_input)
            
        st.session_state.processed = True

if st.session_state.processed:
    st.divider()
    c1, c2 = st.columns([1, 3])
    with c1:
        st.metric("Match Score", f"{st.session_state.analysis['score']}%")
    with c2:
        st.info(f"**Advice:** {st.session_state.analysis['advice']}")

    tab1, tab2 = st.tabs(["📝 Optimized Word Doc", "✉️ Cover Letter"])
    
    with tab1:
        st.success("Resume rewritten. Download the Word Doc to tweak formatting.")
        st.download_button(
            label="📥 Download Resume (.docx)",
            data=st.session_state.docx_data,
            file_name="Optimized_Salesforce_Resume.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
        
    with tab2:
        st.text_area("Cover Letter:", value=st.session_state.cover_letter, height=300)
        st.download_button(
            label="📥 Download Cover Letter",
            data=st.session_state.cover_letter,
            file_name="Cover_Letter.txt"
        )
